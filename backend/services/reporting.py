from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

from backend.schemas import ThreatRecord


def _decision_label(value: str) -> str:
    mapping = {
        "malicious": "zararli",
        "suspicious": "supheli",
        "under_review": "incelemede",
        "benign": "temiz",
    }
    return mapping.get(value, value)


def _severity_label(value: str) -> str:
    mapping = {
        "critical": "kritik",
        "medium": "orta",
        "low": "dusuk",
    }
    return mapping.get(value, value)


def _certainty_bucket(confidence: float) -> str:
    if confidence >= 0.9:
        return "yuksek güven"
    if confidence >= 0.75:
        return "orta-yuksek güven"
    if confidence >= 0.6:
        return "orta güven"
    return "dusuk güven"


def _risk_story(risk_score: int) -> str:
    if risk_score >= 90:
        return "Risk skoru cok yuksek. Izole etme ve erişim kesme aksiyonu gecikmeden uygulanmali."
    if risk_score >= 75:
        return "Risk skoru yuksek. Olayin yatay yayilim ihtimali nedeniyle aktif takip önerilir."
    if risk_score >= 55:
        return "Risk skoru orta. Kontrollu muhafaza ve manuel gozden gecirme birlikte ilerlemeli."
    return "Risk skoru dusuk. Tam engelleme yerine izleme ve ek telemetri toplanmasi uygun goruluyor."


def _fake_ai_thoughts(threat: ThreatRecord) -> list[str]:
    thought_pool = [
        "Dosya uzantisi ile davranış paterni ayni yonde sinyal veriyor, skor arttirildi.",
        "Davranis zincirindeki ilk adim zararsiz gorunse de ikinci adimda niyet değişimi var.",
        "Ayni IP'ye tekrar eden cikis denemesi C2 olasiligini guclendiriyor.",
        "Imza tabanli sinyal tek basina yeterli degil, davranış tabanli kanitla birlestirildi.",
        "Yanlis pozitif riskini azaltmak için belirsizlik notu rapora eklendi.",
        "Muhafaza aksiyonu is surekliligini minimum etkileyecek sekilde önerildi.",
    ]
    if threat.decision == "benign":
        return [
            "Model, zararsiz olasiligini yuksek goruyor ama izleme devam etmeli.",
            "Imza tarafinda güçlü malware sinyali bulunmadi.",
            "Davranis oruntusu kritik esik altinda kaldigi için tam engelleme önerilmedi.",
        ]
    if threat.decision == "under_review":
        return thought_pool[1:5]
    if threat.decision == "suspicious":
        return thought_pool[0:5]
    return thought_pool[:5]


def _counter_points(threat: ThreatRecord) -> list[str]:
    counters = [
        "Dosya mesru bir idari script olabilir; bağlam dogrulamasi gerekli.",
        "IP adresi gecici paylasimli altyapi olabilir; WHOIS + threat intel dogrulanmali.",
        "Tek bir imza eslesmesi, kesin malware karari için yetersiz olabilir.",
        "Izole ortam ile uretilen davranış, uretim ortamindan farkli olasiliklar barindirir.",
    ]
    if threat.risk_score >= 85:
        return counters[:3]
    return counters[1:]


def _model_signals(threat: ThreatRecord) -> list[str]:
    ext = threat.file_name.rsplit(".", maxsplit=1)[-1].lower() if "." in threat.file_name else "bilinmiyor"
    return [
        f"Dosya uzantisi sinyali: .{ext}",
        f"Davranis etiketi: {threat.behavior}",
        f"SHA256 iz: {threat.sha256[:18]}...",
        f"Dosya boyutu: {threat.file_size_bytes} bayt",
    ]


def build_report_payload(
    threat: ThreatRecord,
    blocked_ip: bool,
    latest_alerts: list[str],
    latest_agent_logs: list[str],
) -> dict[str, object]:
    created_at = datetime.now(UTC).isoformat()
    decision_label = _decision_label(threat.decision)
    severity_label = _severity_label(threat.severity)
    confidence_percent = round(threat.confidence * 100)
    certainty = _certainty_bucket(threat.confidence)

    report = {
        "olay_id": f"INC-{threat.id}",
        "başlık": "Agentic SOC Olay Raporu",
        "olüsturulma_zamani": created_at,
        "yonetici_özeti": (
            f"{threat.file_name} dosyasi {decision_label} olarak siniflandi. "
            f"Risk seviyesi {severity_label}, güven %{confidence_percent}."
        ),
        "tehdit_profili": {
            "tehdit_id": threat.id,
            "dosya_adi": threat.file_name,
            "kaynak_ip": threat.source_ip,
            "risk_skoru": threat.risk_score,
            "seviye": severity_label,
            "karar": decision_label,
            "karar_kodu": threat.decision,
            "güven": threat.confidence,
            "sha256": threat.sha256,
            "dosya_boyutu_bayt": threat.file_size_bytes,
            "ip_engelli_mi": blocked_ip,
        },
        "xai_özeti": {
            "ana_gerekceler": threat.reasoning,
            "model_sinyalleri": _model_signals(threat),
            "karsi_gorusler": _counter_points(threat),
            "belirsizlik_notu": f"Model güven seviyesi: {certainty}.",
            "risk_hikayesi": _risk_story(threat.risk_score),
        },
        "ai_dusunce_notlari": _fake_ai_thoughts(threat),
        "önerilen_aksiyonlar": [
            "Kaynak IP'yi gecici olarak karantinaya al.",
            "Ayni hash'e sahip dosyalar için ortam taramasi baslat.",
            "Benzer IOC'ler için son 24 saatlik loglari geriye dönük tara.",
            "Gerekiyorsa olay kaydini hukuk ve uyumluluk birimine ilet.",
        ],
        "operasyon_zaman_cizelgesi": [
            "Telemetri alindi",
            "LLM siniflandirmasi tamamlandi",
            "Sandbox davranışi analiz edildi",
            "Ajanik karar uretildi",
            "SOC paneli guncellendi",
        ],
        "canlı_alarm_özeti": latest_alerts,
        "ajan_gunluk_özeti": latest_agent_logs,
        "uyumluluk_notu": "Veri işleme akisinda dis ortama ham dosya aktarimi yapilmadi.",
        "is_etkisi": (
            "Olay kritik kaynaklara yayilmadan kontrol altina alinirse operasyonel etki dusuk tutulabilir."
        ),
        "kapanis": "Bu rapor demo amaclidir; AI dusunce notlari simulasyon icerir.",
    }
    return report


def build_report_docx(report: dict[str, object]) -> bytes:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    document.add_heading("Agentic SOC Olay Raporu", level=0)
    document.add_paragraph(
        f"Olay No: {report['olay_id']} | Uretim Zamani: {report['olüsturulma_zamani']}"
    )

    document.add_heading("Yonetici Ozeti", level=1)
    document.add_paragraph(str(report["yonetici_özeti"]))

    profile = report["tehdit_profili"]
    document.add_heading("Tehdit Profili", level=1)
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    profile_rows = [
        ("Tehdit ID", str(profile["tehdit_id"])),
        ("Dosya Adi", str(profile["dosya_adi"])),
        ("Kaynak IP", str(profile["kaynak_ip"])),
        ("Risk Skoru", str(profile["risk_skoru"])),
        ("Seviye", str(profile["seviye"])),
        ("Karar", str(profile["karar"])),
        ("Güven", f"%{round(float(profile['güven']) * 100)}"),
        ("SHA256", str(profile["sha256"])),
        ("Dosya Boyutu", f"{profile['dosya_boyutu_bayt']} bayt"),
        ("IP Engelli Mi", "evet" if bool(profile["ip_engelli_mi"]) else "hayir"),
    ]
    for key, value in profile_rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    xai = report["xai_özeti"]
    document.add_heading("Aciklanabilir YZ Ozet", level=1)
    document.add_paragraph(str(xai["risk_hikayesi"]))
    document.add_paragraph(str(xai["belirsizlik_notu"]))

    document.add_heading("Ana Gerekceler", level=2)
    for item in xai["ana_gerekceler"]:
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Model Sinyalleri", level=2)
    for item in xai["model_sinyalleri"]:
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Karsi Gorusler", level=2)
    for item in xai["karsi_gorusler"]:
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("AI Dusunce Notlari (Simulasyon)", level=1)
    for item in report["ai_dusunce_notlari"]:
        document.add_paragraph(str(item), style="List Number")

    document.add_heading("Onerilen Aksiyonlar", level=1)
    for item in report["önerilen_aksiyonlar"]:
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Operasyon Zaman Cizelgesi", level=1)
    for item in report["operasyon_zaman_cizelgesi"]:
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Canlı Alarm Ozeti", level=1)
    for item in report["canlı_alarm_özeti"]:
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Ajan Gunluk Ozeti", level=1)
    for item in report["ajan_gunluk_özeti"]:
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Uyumluluk ve Is Etkisi", level=1)
    document.add_paragraph(str(report["uyumluluk_notu"]))
    document.add_paragraph(str(report["is_etkisi"]))

    document.add_paragraph(str(report["kapanis"]))

    output = BytesIO()
    document.save(output)
    return output.getvalue()
